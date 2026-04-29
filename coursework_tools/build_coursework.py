from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
ASSETS_DIR = ROOT / "coursework_assets"
SCREENSHOTS_DIR = ASSETS_DIR / "screenshots"
OUTPUT_DOCX = ROOT / "Kursovaya_rabota_Biblioteka_2026.docx"
ER_DIAGRAM = ASSETS_DIR / "er_diagram_biblioteka.png"


ACCENT = RGBColor(0x16, 0x5D, 0x59)
ACCENT_DARK = RGBColor(0x0F, 0x17, 0x2A)
TEXT_GRAY = RGBColor(0x44, 0x4F, 0x5A)
LIGHT_BORDER = RGBColor(0xD5, 0xDE, 0xE6)


def ensure_dirs() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        ("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        ("C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
        ("C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf"),
    ]

    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)

    return ImageFont.load_default()


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str],
    header_fill: tuple[int, int, int],
    body_fill: tuple[int, int, int],
) -> None:
    x1, y1, x2, y2 = box
    title_font = load_font(28, bold=True)
    body_font = load_font(22)

    draw.rounded_rectangle([x1, y1, x2, y2], radius=18, fill=body_fill, outline=(190, 202, 214), width=3)
    draw.rounded_rectangle([x1, y1, x2, y1 + 60], radius=18, fill=header_fill)
    draw.rectangle([x1, y1 + 40, x2, y1 + 60], fill=header_fill)
    draw.text((x1 + 18, y1 + 13), title, font=title_font, fill=(255, 255, 255))

    current_y = y1 + 84
    for line in lines:
        draw.text((x1 + 20, current_y), line, font=body_font, fill=(32, 41, 55))
        current_y += 34


def draw_connector(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    left_label: str,
    right_label: str,
) -> None:
    draw.line([start, end], fill=(58, 71, 84), width=4)
    mid_x = (start[0] + end[0]) // 2
    mid_y = (start[1] + end[1]) // 2
    label_font = load_font(22, bold=True)
    draw.text((start[0] + 8, start[1] - 28), left_label, font=label_font, fill=(15, 23, 42))
    draw.text((end[0] - 28, end[1] - 28), right_label, font=label_font, fill=(15, 23, 42))
    draw.ellipse([mid_x - 6, mid_y - 6, mid_x + 6, mid_y + 6], fill=(22, 93, 89))


def build_er_diagram(path: Path) -> None:
    width, height = 1800, 1100
    img = Image.new("RGB", (width, height), (248, 250, 252))
    draw = ImageDraw.Draw(img)

    title_font = load_font(38, bold=True)
    subtitle_font = load_font(22)
    draw.text((70, 40), "ER-диаграмма базы данных «Библиотека»", font=title_font, fill=(15, 23, 42))
    draw.text(
        (70, 95),
        "Концептуальная модель отражает хранение авторов, читателей, изданий, экземпляров книг и операций выдачи.",
        font=subtitle_font,
        fill=(71, 85, 105),
    )

    boxes = {
        "authors": (80, 210, 520, 470),
        "books": (640, 210, 1110, 575),
        "readers": (1290, 210, 1720, 505),
        "copies": (640, 675, 1110, 995),
        "loans": (1290, 655, 1720, 1010),
    }

    draw_box(
        draw,
        boxes["authors"],
        "Authors",
        ["PK AuthorId", "FullName", "UNIQUE FullName"],
        (22, 93, 89),
        (232, 244, 242),
    )
    draw_box(
        draw,
        boxes["books"],
        "Books",
        ["PK BookId", "FK AuthorId", "Title", "Publisher", "PublishYear", "Genre", "ISBN"],
        (15, 23, 42),
        (240, 245, 250),
    )
    draw_box(
        draw,
        boxes["readers"],
        "Readers",
        ["PK ReaderId", "FullName", "Phone", "Address", "TicketNumber", "UNIQUE TicketNumber"],
        (37, 99, 235),
        (239, 246, 255),
    )
    draw_box(
        draw,
        boxes["copies"],
        "BookCopies",
        ["PK CopyId", "FK BookId", "InventoryNumber", "Location", "Status", "UNIQUE InventoryNumber"],
        (180, 83, 9),
        (255, 247, 237),
    )
    draw_box(
        draw,
        boxes["loans"],
        "Loans",
        ["PK LoanId", "FK ReaderId", "FK CopyId", "IssueDate", "DueDate", "ReturnDate", "UX Active Copy"],
        (124, 58, 237),
        (245, 243, 255),
    )

    draw_connector(draw, (520, 340), (640, 340), "1", "N")
    draw_connector(draw, (875, 575), (875, 675), "1", "N")
    draw_connector(draw, (1505, 505), (1505, 655), "1", "N")
    draw_connector(draw, (1110, 830), (1290, 830), "1", "N")

    note_font = load_font(20)
    draw.rounded_rectangle([80, 880, 520, 1015], radius=16, fill=(255, 255, 255), outline=(190, 202, 214), width=2)
    note_lines = [
        "Бизнес-правила:",
        "1) выдать можно только доступный экземпляр;",
        "2) у одного экземпляра не более одной активной выдачи;",
        "3) дата возврата не может быть раньше даты выдачи.",
    ]
    y = 905
    for line in note_lines:
        draw.text((100, y), line, font=note_font, fill=(51, 65, 85))
        y += 26

    img.save(path)


def apply_font(run, size: float = 14, bold: bool = False, italic: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_page_margins(section) -> None:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.different_first_page_header_footer = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    apply_font(run, size=12)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    set_page_margins(section)
    add_page_number(section.footer.paragraphs[0])

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(14)


def body_paragraph(doc: Document, text: str, first_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    apply_font(run, size=14)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        size = 16
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        size = 15
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        size = 14
    run = p.add_run(text)
    apply_font(run, size=size, bold=True)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    apply_font(run, size=14)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    apply_font(run, size=12, italic=True)


def add_image(doc: Document, image_path: Path, width_cm: float, caption_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    caption(doc, caption_text)


def add_table(doc: Document, rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False

    for idx, width in enumerate(widths_cm):
        table.columns[idx].width = Cm(width)

    header_cells = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        header_cells[idx].text = value
        for paragraph in header_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                apply_font(run, size=12, bold=True)

    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    apply_font(run, size=12)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(10)


def add_title_page(doc: Document) -> None:
    items = [
        "Министерство образования Саратовской области",
        "Государственное автономное профессиональное образовательное учреждение Саратовской области",
        "«Новоузенский агротехнологический техникум»",
        "",
        "",
        "Курсовая работа",
        "по ПМ.04 «Разработка, администрирование и защита баз данных»",
        "",
        "Тема: Разработка базы данных и приложения «Библиотека»",
        "",
        "",
        "Выполнил: студент ________________________________",
        "Проверил: _______________________________________",
        "",
        "",
        "г. Новоузенск, 2026 год",
    ]

    for idx, text in enumerate(items):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if idx in {5, 8}:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
        if idx in {11, 12}:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        apply_font(
            run,
            size=16 if idx == 5 else 14,
            bold=idx in {5, 8},
            color=ACCENT_DARK if idx in {5, 8} else None,
        )

    doc.add_page_break()


def build_document() -> None:
    ensure_dirs()
    build_er_diagram(ER_DIAGRAM)

    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    heading(doc, "ВВЕДЕНИЕ", level=1)
    body_paragraph(
        doc,
        "Современная библиотека работает с большим количеством взаимосвязанных данных: читателями, "
        "авторами, изданиями, конкретными экземплярами книг, фактами выдачи и возврата. При ведении "
        "учета вручную библиотекарь тратит много времени на поиск карточек, контроль доступности книг, "
        "проверку задолженности и подготовку отчетов. Это повышает вероятность ошибок, дублирования "
        "записей и потери актуальной информации."
    )
    body_paragraph(
        doc,
        "Актуальность работы состоит в необходимости автоматизации библиотечного фонда на основе "
        "реляционной базы данных и удобного прикладного интерфейса. Программное решение должно "
        "обеспечивать централизованное хранение сведений, ускорять типовые операции и повышать "
        "достоверность учета."
    )
    body_paragraph(
        doc,
        "Цель курсовой работы — разработать базу данных и настольное приложение «Библиотека», "
        "предназначенные для учета читателей, книг, экземпляров, выдач, возвратов и формирования "
        "оперативной отчетности."
    )
    body_paragraph(doc, "Для достижения поставленной цели необходимо решить следующие задачи:")
    bullet(doc, "изучить предметную область учета библиотечного фонда и выделить основные сущности;")
    bullet(doc, "спроектировать концептуальную, логическую и физическую модели базы данных;")
    bullet(doc, "реализовать структуру базы данных в SQL Server и заполнить ее демонстрационными данными;")
    bullet(doc, "подготовить прикладное WPF-приложение для работы с данными библиотеки;")
    bullet(doc, "проанализировать основные пользовательские сценарии и экранные формы программы.")

    heading(doc, "1. Теоретические основы построения базы данных «Библиотека»", level=1)
    heading(doc, "1.1 Постановка задачи", level=2)
    body_paragraph(
        doc,
        "Библиотечный учет включает операции регистрации читателей, ведения каталога изданий, "
        "фиксации местонахождения конкретных экземпляров, оформления выдач и возвратов, а также "
        "контроля просроченных операций. Основная задача автоматизации заключается в создании "
        "целостной информационной системы, которая исключает противоречия данных и обеспечивает "
        "быстрый доступ к актуальному состоянию фонда."
    )
    body_paragraph(
        doc,
        "Разрабатываемая система должна поддерживать добавление, изменение и удаление записей, "
        "поиск по ключевым полям, фильтрацию по статусу экземпляров, формирование каталожной карточки "
        "выбранной книги и вывод отчетов по фонду. Отдельное внимание уделяется проверке бизнес-правил: "
        "выдать можно только доступный экземпляр, а один экземпляр не должен иметь две активные выдачи."
    )

    heading(doc, "1.2 Анализ предметной области", level=2)
    body_paragraph(
        doc,
        "Предметная область охватывает деятельность библиотекаря по обслуживанию читателей и учету "
        "движения книг. Информационные объекты системы делятся на постоянные и оперативные. "
        "К постоянным относятся авторы, книги, читатели и экземпляры книг. К оперативным относятся "
        "записи о выдачах, в которых фиксируются даты выдачи, планового возврата и фактического возврата."
    )
    body_paragraph(doc, "В рамках анализа предметной области выделены следующие основные процессы:")
    bullet(doc, "ведение справочника читателей с хранением ФИО, номера читательского билета, телефона и адреса;")
    bullet(doc, "ведение списка авторов и каталога изданий с привязкой каждой книги к автору;")
    bullet(doc, "учет конкретных экземпляров книг по инвентарному номеру, месту хранения и статусу;")
    bullet(doc, "оформление выдачи книги читателю с контролем доступности экземпляра;")
    bullet(doc, "оформление возврата книги и автоматическое изменение статуса экземпляра;")
    bullet(doc, "формирование списка должников и аналитических отчетов по библиотечному фонду.")
    body_paragraph(
        doc,
        "Анализ действующих процессов показывает, что учет целесообразно организовать на основе "
        "реляционной базы данных. Такой подход позволяет формализовать связи между сущностями, "
        "обеспечить ссылочную целостность и сократить число ошибок при обработке повседневных операций."
    )

    heading(doc, "1.3 Проектирование базы данных", level=2)
    heading(doc, "1.3.1 Концептуальная модель (ER-диаграмма)", level=3)
    body_paragraph(
        doc,
        "Концептуальная модель базы данных «Библиотека» включает пять ключевых сущностей: Authors, "
        "Books, Readers, BookCopies и Loans. Сущность Authors хранит сведения об авторах. Сущность Books "
        "описывает сами издания и связывает каждую книгу с автором. Сущность BookCopies отражает "
        "конкретные экземпляры книг, доступные в фонде. Сущность Readers хранит информацию о пользователях "
        "библиотеки. Сущность Loans описывает факты выдачи экземпляров читателям."
    )
    add_image(doc, ER_DIAGRAM, 16.5, "Рис. 1. ER-диаграмма базы данных «Библиотека»")

    heading(doc, "1.3.2 Логическая модель (таблицы)", level=3)
    body_paragraph(
        doc,
        "Логическая модель представлена пятью взаимосвязанными таблицами. Для каждой сущности "
        "выбран первичный ключ типа INT IDENTITY, а связи между таблицами оформлены внешними ключами. "
        "Дополнительно введены уникальные ограничения для номера читательского билета, полного имени автора "
        "и инвентарного номера экземпляра."
    )
    add_table(
        doc,
        [
            ["Таблица", "Назначение", "Ключевые поля"],
            ["Authors", "Справочник авторов", "AuthorId (PK), FullName"],
            ["Readers", "Карточки читателей", "ReaderId (PK), TicketNumber (UQ)"],
            ["Books", "Каталог изданий", "BookId (PK), AuthorId (FK), Title"],
            ["BookCopies", "Учет экземпляров", "CopyId (PK), BookId (FK), InventoryNumber (UQ)"],
            ["Loans", "Операции выдачи", "LoanId (PK), ReaderId (FK), CopyId (FK)"],
        ],
        [3.2, 6.6, 7.2],
    )

    heading(doc, "1.3.3 Физическая модель", level=3)
    body_paragraph(
        doc,
        "Физическая модель реализована в СУБД SQL Server Express. В таблицах используются типы "
        "INT IDENTITY, NVARCHAR и DATE. Для ускорения поиска предусмотрены индексы по названию книги, "
        "ФИО автора, статусу экземпляра и сроку возврата. Дополнительный фильтрованный уникальный индекс "
        "на таблице Loans исключает наличие двух одновременно активных выдач для одного экземпляра книги."
    )

    heading(doc, "2. Реализация базы данных", level=1)
    heading(doc, "2.1 SQL-код создания структуры базы данных", level=2)
    body_paragraph(
        doc,
        "Создание структуры базы данных выполняется SQL-скриптом LibraryDB.sql. Скрипт формирует "
        "таблицы Readers, Authors, Books, BookCopies и Loans, определяет первичные и внешние ключи, "
        "проверочные ограничения и индексы."
    )
    add_code_block(
        doc,
        [
            "CREATE TABLE dbo.Readers (",
            "    ReaderId INT IDENTITY(1,1) NOT NULL PRIMARY KEY,",
            "    FullName NVARCHAR(100) NOT NULL,",
            "    Phone NVARCHAR(20) NULL,",
            "    Address NVARCHAR(200) NULL,",
            "    TicketNumber NVARCHAR(30) NOT NULL UNIQUE",
            ");",
            "",
            "CREATE TABLE dbo.Authors (",
            "    AuthorId INT IDENTITY(1,1) NOT NULL PRIMARY KEY,",
            "    FullName NVARCHAR(100) NOT NULL UNIQUE",
            ");",
            "",
            "CREATE TABLE dbo.Books (",
            "    BookId INT IDENTITY(1,1) NOT NULL PRIMARY KEY,",
            "    Title NVARCHAR(150) NOT NULL,",
            "    AuthorId INT NOT NULL,",
            "    Publisher NVARCHAR(100) NULL,",
            "    PublishYear INT NULL,",
            "    Genre NVARCHAR(50) NULL,",
            "    ISBN NVARCHAR(30) NULL,",
            "    CONSTRAINT FK_Books_Authors FOREIGN KEY (AuthorId) REFERENCES dbo.Authors(AuthorId)",
            ");",
        ],
    )
    body_paragraph(
        doc,
        "Для таблицы BookCopies предусмотрен признак статуса экземпляра. Допустимыми значениями статуса "
        "являются «доступна», «выдана» и «потеряна». Для таблицы Loans установлено ограничение на корректность "
        "дат и фильтрованный уникальный индекс, запрещающий создание второй активной выдачи для того же экземпляра."
    )

    heading(doc, "2.2 Заполнение базы демонстрационными данными", level=2)
    body_paragraph(
        doc,
        "Скрипт LibraryDB.Seed.sql наполняет базу тестовыми данными, достаточными для демонстрации "
        "работы приложения. На момент подготовки курсовой в базе присутствуют 12 читателей, 10 авторов, "
        "12 изданий, 24 экземпляра книг и 11 записей о выдаче. Наличие заполненной базы позволяет "
        "проверить основные сценарии работы: поиск, регистрацию новых объектов, оформление выдачи и возврата, "
        "а также формирование отчетов."
    )
    body_paragraph(
        doc,
        "Демонстрационное наполнение охватывает как завершенные, так и активные выдачи, а также экземпляры "
        "с различными статусами. Благодаря этому можно корректно проверить список должников, работу фильтров "
        "по статусу и отображение показателей на главной панели приложения."
    )

    heading(doc, "2.3 Примеры запросов к базе данных", level=2)
    body_paragraph(
        doc,
        "В прикладной части используются параметризованные запросы, реализованные в классе "
        "LibraryRepository. Ниже приведены характерные примеры запросов, отражающие поиск книг, "
        "получение активных выдач и вычисление показателей для рабочей панели."
    )
    add_code_block(
        doc,
        [
            "SELECT b.BookId, b.Title, a.FullName AS AuthorName,",
            "       b.Publisher, b.PublishYear, b.Genre, b.ISBN,",
            "       COUNT(c.CopyId) AS CopyCount",
            "FROM Books b",
            "INNER JOIN Authors a ON a.AuthorId = b.AuthorId",
            "LEFT JOIN BookCopies c ON c.BookId = b.BookId",
            "WHERE (@title = N'' OR b.Title LIKE N'%' + @title + N'%')",
            "GROUP BY b.BookId, b.Title, a.FullName, b.Publisher, b.PublishYear, b.Genre, b.ISBN;",
            "",
            "SELECT l.LoanId, r.FullName AS ReaderName, b.Title AS BookTitle,",
            "       c.InventoryNumber, l.IssueDate, l.DueDate",
            "FROM Loans l",
            "INNER JOIN Readers r ON r.ReaderId = l.ReaderId",
            "INNER JOIN BookCopies c ON c.CopyId = l.CopyId",
            "INNER JOIN Books b ON b.BookId = c.BookId",
            "WHERE l.ReturnDate IS NULL",
            "ORDER BY l.DueDate, r.FullName;",
        ],
    )

    heading(doc, "2.4 Описание сущностей и атрибутов базы данных", level=2)
    heading(doc, "2.4.1 Сущности и их атрибуты", level=3)
    body_paragraph(doc, "Ниже приведена краткая характеристика основных сущностей базы данных.")
    bullet(doc, "Authors: код автора и полное имя автора. Таблица используется как справочник.")
    bullet(doc, "Readers: код читателя, ФИО, телефон, адрес и номер читательского билета.")
    bullet(doc, "Books: код книги, автор, название, издательство, год издания, жанр и ISBN.")
    bullet(doc, "BookCopies: код экземпляра, ссылка на книгу, инвентарный номер, место хранения и статус.")
    bullet(doc, "Loans: код выдачи, ссылки на читателя и экземпляр, дата выдачи, срок возврата и дата возврата.")

    heading(doc, "2.4.2 Связи между сущностями", level=3)
    body_paragraph(
        doc,
        "Связь между таблицами Authors и Books имеет тип «один ко многим»: одному автору может "
        "соответствовать несколько книг. Аналогично таблица Books связана с таблицей BookCopies: "
        "одна книга может иметь несколько экземпляров. Таблица Readers связана с таблицей Loans "
        "отношением «один ко многим», поскольку один читатель может получить несколько книг. "
        "Таблица BookCopies также связана с таблицей Loans отношением «один ко многим», однако "
        "с учетом фильтрованного индекса одновременно допускается не более одной активной выдачи."
    )

    heading(doc, "2.4.3 Бизнес-правила и ограничения целостности", level=3)
    body_paragraph(doc, "Для обеспечения корректности работы системы введены следующие ограничения:")
    bullet(doc, "ФИО читателя и номер читательского билета обязательны при сохранении карточки читателя.")
    bullet(doc, "Название книги и выбранный автор обязательны при создании записи об издании.")
    bullet(doc, "Инвентарный номер экземпляра должен быть уникальным.")
    bullet(doc, "Выдать можно только экземпляр со статусом «доступна».")
    bullet(doc, "Срок возврата не может быть меньше даты выдачи, а дата возврата — меньше даты выдачи.")
    bullet(doc, "Удаление записи ограничивается ссылочной целостностью: связанные данные должны учитываться.")

    heading(doc, "3. Создание приложения для базы данных «Библиотека»", level=1)
    heading(doc, "3.1 Общая характеристика приложения", level=2)
    body_paragraph(
        doc,
        "Для работы с базой данных разработано настольное приложение на платформе .NET 8 с использованием "
        "технологии WPF. Приложение реализовано в виде одного главного окна с системой вкладок. "
        "Такой подход позволяет библиотекарю быстро переключаться между справочниками, карточками, "
        "операциями выдачи и отчетами без открытия дополнительных форм."
    )
    body_paragraph(
        doc,
        "Взаимодействие с базой данных инкапсулировано в классе LibraryRepository. В нем реализованы "
        "запросы на чтение и изменение данных, а также проверки бизнес-правил. Пользовательский интерфейс "
        "связан с коллекциями объектов через механизм привязки данных WPF, что упрощает обновление экранных форм."
    )

    heading(doc, "3.2 Описание интерфейса и основных экранных форм", level=2)
    body_paragraph(
        doc,
        "Главная вкладка приложения выполняет роль рабочей панели библиотекаря. На ней отображаются "
        "ключевые показатели: число читателей, изданий, экземпляров, доступных книг, активных выдач и должников. "
        "Также предусмотрены кнопки быстрого перехода к основным разделам системы."
    )
    add_image(doc, SCREENSHOTS_DIR / "01-dashboard.png", 16.5, "Рис. 2. Главное окно приложения «Библиотека»")
    body_paragraph(
        doc,
        "Раздел «Читатели» предназначен для поиска карточек по ФИО и номеру билета, просмотра списка "
        "читателей и ввода новых записей. В правой части окна расположена форма редактирования с обязательными "
        "полями ФИО и номера читательского билета."
    )
    add_image(doc, SCREENSHOTS_DIR / "02-readers.png", 16.5, "Рис. 3. Раздел «Читатели»")
    body_paragraph(
        doc,
        "Раздел «Издания» используется для ведения книжного каталога. Пользователь может выполнять поиск "
        "по названию, автору, жанру и году, а также добавлять новых авторов прямо из формы карточки издания. "
        "В таблице отображается число экземпляров каждой книги."
    )
    add_image(doc, SCREENSHOTS_DIR / "03-books.png", 16.5, "Рис. 4. Раздел «Издания»")
    body_paragraph(
        doc,
        "Раздел «Экземпляры» хранит сведения о физических экземплярах книг, включая их инвентарные номера, "
        "местонахождение и текущий статус. Именно этот раздел используется для оперативного контроля фонда."
    )
    add_image(doc, SCREENSHOTS_DIR / "04-copies.png", 16.5, "Рис. 5. Раздел «Экземпляры»")

    heading(doc, "3.3 Основные пользовательские сценарии", level=2)
    body_paragraph(
        doc,
        "Оформление выдачи реализовано в отдельной вкладке. Библиотекарь выбирает читателя, доступный "
        "экземпляр, дату выдачи и срок возврата. В правой части экрана одновременно выводится список активных "
        "выдач, что позволяет контролировать текущую загрузку фонда."
    )
    add_image(doc, SCREENSHOTS_DIR / "05-issue.png", 16.5, "Рис. 6. Оформление выдачи книги")
    body_paragraph(
        doc,
        "Процесс возврата также вынесен в самостоятельную вкладку. После выбора записи о выдаче система "
        "обновляет дату возврата и автоматически переводит экземпляр в статус «доступна». Такой сценарий "
        "упрощает контроль движения фонда и исключает ручное изменение нескольких таблиц."
    )
    add_image(doc, SCREENSHOTS_DIR / "06-return.png", 16.5, "Рис. 7. Оформление возврата книги")
    body_paragraph(
        doc,
        "Раздел «Отчеты» предоставляет четыре аналитических представления: список всех книг, доступных "
        "экземпляров, выданных книг и должников. Это позволяет использовать систему не только для учета, "
        "но и для быстрого получения управленческой информации."
    )
    add_image(doc, SCREENSHOTS_DIR / "07-reports.png", 16.5, "Рис. 8. Раздел «Отчеты»")
    body_paragraph(
        doc,
        "В результате приложение охватывает полный цикл повседневной работы библиотекаря: от создания "
        "карточек и наполнения каталога до оформления выдачи, возврата и анализа состояния фонда."
    )

    heading(doc, "ЗАКЛЮЧЕНИЕ", level=1)
    body_paragraph(
        doc,
        "В ходе выполнения курсовой работы была разработана база данных и настольное приложение "
        "«Библиотека», предназначенные для автоматизации учета библиотечного фонда. В процессе работы "
        "были проанализированы особенности предметной области, спроектированы концептуальная, логическая "
        "и физическая модели данных, а затем реализована структура базы данных в SQL Server."
    )
    body_paragraph(
        doc,
        "Программная реализация на WPF обеспечивает удобный доступ к справочникам, экземплярам, выдачам, "
        "возвратам и отчетам. Использование ограничений целостности и бизнес-правил повышает надежность "
        "системы и снижает вероятность ошибок при работе библиотекаря."
    )
    body_paragraph(
        doc,
        "Практическая значимость работы заключается в возможности непосредственного применения разработанной "
        "системы в учебной или небольшой библиотеке. Дальнейшее развитие проекта может быть связано с "
        "введением авторизации пользователей, печатных форм, резервного копирования и расширенной аналитики."
    )

    heading(doc, "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ", level=1)
    references = [
        "Дейт К. Дж. Введение в системы баз данных. — М.: Вильямс, 2020.",
        "Кузнецов С. Д. Основы баз данных. — М.: Интернет-Университет Информационных Технологий, 2021.",
        "Голицына О. Л., Максимов Н. В., Попов И. И. Базы данных. — М.: Форум, 2022.",
        "Microsoft. SQL Server documentation. — URL: https://learn.microsoft.com/sql/",
        "Microsoft. WPF documentation. — URL: https://learn.microsoft.com/dotnet/desktop/wpf/",
        "Microsoft. .NET documentation. — URL: https://learn.microsoft.com/dotnet/",
        "Исходный код проекта Biblioteka: файлы MainWindow.xaml, MainWindow.xaml.cs, LibraryRepository.cs, LibraryDB.sql.",
        "Скрипт демонстрационного наполнения базы данных LibraryDB.Seed.sql.",
    ]
    for index, item in enumerate(references, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(f"{index}. {item}")
        apply_font(run, size=14)

    doc.add_page_break()
    heading(doc, "ПРИЛОЖЕНИЕ", level=1)
    heading(doc, "Фрагменты SQL-скрипта создания базы данных", level=2)
    add_code_block(
        doc,
        [
            "CREATE TABLE dbo.BookCopies (",
            "    CopyId INT IDENTITY(1,1) NOT NULL PRIMARY KEY,",
            "    BookId INT NOT NULL,",
            "    InventoryNumber NVARCHAR(30) NOT NULL UNIQUE,",
            "    Location NVARCHAR(100) NOT NULL,",
            "    Status NVARCHAR(30) NOT NULL DEFAULT N'доступна',",
            "    CONSTRAINT FK_BookCopies_Books FOREIGN KEY (BookId) REFERENCES dbo.Books(BookId),",
            "    CONSTRAINT CK_BookCopies_Status CHECK (Status IN (N'доступна', N'выдана', N'потеряна'))",
            ");",
            "",
            "CREATE TABLE dbo.Loans (",
            "    LoanId INT IDENTITY(1,1) NOT NULL PRIMARY KEY,",
            "    ReaderId INT NOT NULL,",
            "    CopyId INT NOT NULL,",
            "    IssueDate DATE NOT NULL,",
            "    DueDate DATE NOT NULL,",
            "    ReturnDate DATE NULL,",
            "    CONSTRAINT FK_Loans_Readers FOREIGN KEY (ReaderId) REFERENCES dbo.Readers(ReaderId),",
            "    CONSTRAINT FK_Loans_BookCopies FOREIGN KEY (CopyId) REFERENCES dbo.BookCopies(CopyId),",
            "    CONSTRAINT CK_Loans_Dates CHECK (DueDate >= IssueDate)",
            ");",
            "",
            "CREATE UNIQUE INDEX UX_Loans_ActiveCopy",
            "    ON dbo.Loans (CopyId)",
            "    WHERE ReturnDate IS NULL;",
            "",
            "CREATE INDEX IX_Books_Title ON dbo.Books (Title);",
            "CREATE INDEX IX_Authors_FullName ON dbo.Authors (FullName);",
            "CREATE INDEX IX_BookCopies_Status ON dbo.BookCopies (Status);",
            "CREATE INDEX IX_Loans_DueDate ON dbo.Loans (DueDate);",
        ],
    )

    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
